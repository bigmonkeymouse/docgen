#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word模板填充工具
根据Excel数据替换Word模板中的{{占位符}}
支持格式保留和页眉页脚脚注尾注处理
"""

import argparse
import sys
import re
import zipfile
import os
from xml.etree import ElementTree as ET
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from openpyxl import load_workbook
import datetime
import json
from pypinyin import pinyin, Style


def load_excel_data(excel_path):
    """
    从Excel文件加载数据
    新格式：两列，第一列是字段名，第二列是对应的值
    返回字典列表，每行数据为一个字典
    """
    # 验证文件存在性
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel文件不存在: {excel_path}")
    
    workbook = load_workbook(excel_path, data_only=True)
    sheet = workbook.active

    # 读取两列数据
    raw_data = {}
    # 收集股东信息
    shareholders = []
    
    for row in sheet.iter_rows(min_row=1, values_only=True):
        if len(row) >= 2 and row[0] is not None:
            key = str(row[0]).strip()
            value = str(row[1]) if row[1] is not None else ""
            if key:  # 只处理有字段名的行
                raw_data[key] = value
                # 收集股东信息
                if "股东" in key and value:
                    shareholders.append(value)

    # 生成处理后的数据
    processed_data = {}
    
    # 1. 所在分行缩写
    if "所在分行" in raw_data:
        branch = raw_data["所在分行"]
        # 提取城市名（去掉"分行"二字）
        city_name = branch.replace("分行", "")
        # 提取拼音首字母大写
        pinyin_list = pinyin(city_name, style=Style.FIRST_LETTER)
        pinyin_abbr = "".join([p[0].upper() for p in pinyin_list])
        processed_data["所在分行缩写"] = f"{pinyin_abbr}分行"
    
    # 2. 批复金额
    if "批复金额" in raw_data:
        try:
            amount = float(raw_data["批复金额"])
            # 格式化保留两位小数（国际计数法，3位一节用逗号隔开）
            formatted_amount = "{:,.2f}".format(amount)
            # 存储格式化后的金额
            processed_data["批复金额"] = formatted_amount
            # 转换为银行标准大写
            processed_data["批复金额大写"] = num_to_Chinese(amount)
        except ValueError:
            processed_data["批复金额"] = ""
            processed_data["批复金额大写"] = ""
    
    # 3. 额度启用日期和到期日期
    if "额度启用日期" in raw_data:
        enable_date_str = raw_data["额度启用日期"]
        try:
            # 解析日期（格式如"2026年2月"）
            parts = enable_date_str.split("年")
            year = int(parts[0])
            month = int(parts[1].replace("月", ""))
            # 额度启用日期缩写
            processed_data["额度启用日期缩写"] = f"{year}.{month:02d}"
            # 额度到期日期（加一年）
            processed_data["额度到期日期"] = f"{year+1}年{month}月"
        except (IndexError, ValueError):
            processed_data["额度启用日期缩写"] = ""
            processed_data["额度到期日期"] = ""
    
    # 4. 合同制作日期
    current_date = datetime.datetime.now()
    processed_data["合同制作日期"] = current_date.strftime("%Y%m%d")
    
    # 5. 合同制作号
    processed_data["合同制作号"] = generate_contract_number()
    
    # 6. 各类合同编号
    if "所在分行缩写" in processed_data and "部门" in raw_data:
        branch_abbr = processed_data["所在分行缩写"]
        department = raw_data["部门"]
        contract_date = processed_data["合同制作日期"]
        contract_number = processed_data["合同制作号"]
        
        processed_data["综字1"] = f"平银{branch_abbr}{department}综字{contract_date}{contract_number}"
        processed_data["资池字2"] = f"平银{branch_abbr}{department}资池字{contract_date}{contract_number}"
        processed_data["资池质字3"] = f"平银{branch_abbr}{department}资池质字{contract_date}{contract_number}"
        processed_data["自由票字4"] = f"平银{branch_abbr}{department}自由票字{contract_date}{contract_number}"
        processed_data["线融字5"] = f"平银{branch_abbr}{department}线融字{contract_date}{contract_number}"
        processed_data["国内信字6"] = f"平银{branch_abbr}{department}国内信字{contract_date}{contract_number}"
        processed_data["国内信商字7"] = f"平银{branch_abbr}{department}国内信商字{contract_date}{contract_number}"
    
    # 7. 股东信息
    processed_data["股东列表"] = shareholders
    if shareholders:
        if len(shareholders) == 1:
            processed_data["股东"] = shareholders[0]
        else:
            # 多个股东用顿号拼接
            processed_data["股东"] = "、".join(shareholders)
    
    # 合并原始数据和处理后的数据
    data = {**raw_data, **processed_data}
    
    print(f"处理后的数据: {list(processed_data.keys())}")
    print(f"股东信息: {shareholders}")
    
    # 返回包含一个字典的列表
    return [data] if data else []


def num_to_Chinese(num):
    """
    将数字转换为银行标准大写金额
    """
    digits = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
    units = ["", "拾", "佰", "仟"]
    big_units = ["", "万", "亿"]
    
    if num == 0:
        return "零元整"
    
    # 处理整数部分和小数部分
    integer_part = int(num)
    decimal_part = int(round((num - integer_part) * 100))
    
    result = ""
    
    # 处理整数部分
    if integer_part > 0:
        unit_index = 0
        big_unit_index = 0
        temp = ""
        
        while integer_part > 0:
            digit = integer_part % 10
            if digit > 0:
                temp = digits[digit] + units[unit_index] + temp
            else:
                # 避免连续的零
                if temp and not temp.startswith("零"):
                    temp = "零" + temp
            
            unit_index += 1
            if unit_index == 4:
                # 每四位处理一次
                if temp:
                    result = temp + big_units[big_unit_index] + result
                temp = ""
                unit_index = 0
                big_unit_index += 1
            
            integer_part = integer_part // 10
        
        if temp:
            result = temp + big_units[big_unit_index] + result
        
        result += "元"
    
    # 处理小数部分
    if decimal_part == 0:
        result += "整"
    else:
        jiao = decimal_part // 10
        fen = decimal_part % 10
        if jiao > 0:
            result += digits[jiao] + "角"
        if fen > 0:
            result += digits[fen] + "分"
    
    return result


def generate_contract_number():
    """
    生成合同制作号
    规则：每日早上8点开始编号，初始值为"第001号"，每1分30秒递增1，次日早上8点重置
    """
    now = datetime.datetime.now()
    today = now.date()
    
    # 检查是否需要重置
    if now.hour < 8:
        # 凌晨到8点，使用前一天的编号
        yesterday = today - datetime.timedelta(days=1)
        reset_time = datetime.datetime.combine(yesterday, datetime.time(8, 0, 0))
    else:
        reset_time = datetime.datetime.combine(today, datetime.time(8, 0, 0))
    
    # 计算从重置时间到现在经过的秒数
    seconds_since_reset = (now - reset_time).total_seconds()
    
    # 每1分30秒（90秒）递增一个编号
    interval = int(seconds_since_reset // 90)
    
    # 计算应该的编号（从1开始）
    current_number = interval + 1
    
    # 生成编号
    return f"第{current_number:03d}号"


def find_placeholders(text):
    """
    查找文本中的所有占位符
    返回：[(占位符全文本如{{name}}, 占位符名称如name), ...]
    """
    pattern = r'\{\{([^}]+)\}\}'
    matches = re.finditer(pattern, text)
    return [(match.group(0), match.group(1).strip()) for match in matches]


def replace_in_paragraph(paragraph, replacement_dict, debug=False):
    """
    在段落中替换占位符，保留原有格式
    支持占位符跨越多个run的情况
    """
    if not paragraph.runs:
        return False

    # 合并所有run的文本用于匹配
    full_text = "".join([run.text for run in paragraph.runs])

    # 查找所有占位符
    placeholders = find_placeholders(full_text)

    if not placeholders:
        return False

    replaced = False
    if debug:
        print(f"  找到段落: {full_text[:50]}...")

    # 逐个处理每个占位符
    for placeholder_text, placeholder_name in placeholders:
        if placeholder_name in replacement_dict:
            replacement_value = str(replacement_dict[placeholder_name])

            # 在完整文本中找到占位符的位置
            placeholder_pos = full_text.find(placeholder_text)
            if placeholder_pos == -1:
                continue

            placeholder_end = placeholder_pos + len(placeholder_text)

            # 找到占位符跨越的所有run
            current_pos = 0
            runs_to_modify = []
            for run in paragraph.runs:
                run_start = current_pos
                run_end = current_pos + len(run.text)
                current_pos = run_end

                # 检查这个run是否与占位符有交集
                if run_start < placeholder_end and run_end > placeholder_pos:
                    runs_to_modify.append((run, run_start, run_end))

            if runs_to_modify:
                # 检查占位符是否完全在一个run中
                first_run, first_start, first_end = runs_to_modify[0]
                last_run, last_start, last_end = runs_to_modify[-1]

                if len(runs_to_modify) == 1:
                    # 占位符完全在一个run中，直接替换
                    first_run.text = first_run.text.replace(placeholder_text, replacement_value)
                else:
                    # 占位符跨越多个run，需要合并处理
                    # 计算占位符在第一个run中的偏移
                    offset_in_first = placeholder_pos - first_start
                    # 计算占位符在最后一个run中的结束位置
                    offset_in_last = placeholder_end - last_start

                    # 保留第一个run中占位符之前的文本
                    prefix = first_run.text[:offset_in_first] if offset_in_first > 0 else ""
                    # 保留最后一个run中占位符之后的文本
                    suffix = last_run.text[offset_in_last:] if offset_in_last < len(last_run.text) else ""

                    # 清空中间的所有run
                    for run, _, _ in runs_to_modify[1:-1]:
                        run.text = ""

                    # 在第一个run中设置完整替换值
                    first_run.text = prefix + replacement_value + suffix
                    # 清空最后一个run（内容已合并到第一个run）
                    last_run.text = ""

                replaced = True

                if debug:
                    print(f"    替换: {placeholder_text} -> {replacement_value}")

                # 更新full_text用于后续查找
                full_text = "".join([run.text for run in paragraph.runs])

    return replaced


def replace_in_table(table, replacement_dict, debug=False):
    """
    在表格中替换占位符
    """
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, replacement_dict, debug):
                    replaced = True
    return replaced


def replace_in_headers(doc, replacement_dict, debug=False):
    """
    在页眉中替换占位符
    """
    replaced = False

    # 遍历所有section的header
    for i, section in enumerate(doc.sections):
        if section.header:
            if debug:
                print(f"\n处理页眉 Section {i}:")
            for paragraph in section.header.paragraphs:
                if replace_in_paragraph(paragraph, replacement_dict, debug):
                    replaced = True

            # 处理页眉中的表格
            for table in section.header.tables:
                if replace_in_table(table, replacement_dict, debug):
                    replaced = True

    return replaced


def replace_in_footers(doc, replacement_dict, debug=False):
    """
    在页脚中替换占位符
    """
    replaced = False

    # 遍历所有section的footer
    for i, section in enumerate(doc.sections):
        if section.footer:
            if debug:
                print(f"\n处理页脚 Section {i}:")
            for paragraph in section.footer.paragraphs:
                if replace_in_paragraph(paragraph, replacement_dict, debug):
                    replaced = True

            # 处理页脚中的表格
            for table in section.footer.tables:
                if replace_in_table(table, replacement_dict, debug):
                    replaced = True

    return replaced


def replace_in_footers_xml(docx_path, replacement_dict, output_path, debug=False):
    """
    直接操作XML文件，替换所有footer中的占位符
    python-docx可能无法读取某些footer，使用XML直接处理
    """
    replaced = False

    # 创建临时目录
    temp_dir = Path(output_path).parent / "temp"
    temp_dir.mkdir(exist_ok=True)

    # 解压docx文件
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    # 处理所有footer XML文件
    footer_files = list(temp_dir.glob("word/footer*.xml"))

    if debug:
        print(f"\n=== 处理Footer XML文件 ===")
        print(f"找到 {len(footer_files)} 个footer文件")

    for footer_file in footer_files:
        if debug:
            print(f"\n处理: {footer_file.name}")

        # 读取XML
        tree = ET.parse(footer_file)
        root = tree.getroot()

        # 递归替换所有文本节点中的占位符
        def replace_text_in_element(element):
            nonlocal replaced
            if element.text:
                placeholders = find_placeholders(element.text)
                for placeholder_text, placeholder_name in placeholders:
                    if placeholder_name in replacement_dict:
                        replacement_value = str(replacement_dict[placeholder_name])
                        if debug:
                            print(f"  替换: {placeholder_text} -> {replacement_value}")
                        element.text = element.text.replace(placeholder_text, replacement_value)
                        replaced = True

            for child in element:
                replace_text_in_element(child)
                if child.tail:
                    placeholders = find_placeholders(child.tail)
                    for placeholder_text, placeholder_name in placeholders:
                        if placeholder_name in replacement_dict:
                            replacement_value = str(replacement_dict[placeholder_name])
                            if debug:
                                print(f"  替换(尾文本): {placeholder_text} -> {replacement_value}")
                            child.tail = child.tail.replace(placeholder_text, replacement_value)
                            replaced = True

        replace_text_in_element(root)

        # 保存修改后的XML
        tree.write(footer_file, encoding='utf-8', xml_declaration=True)

    # 重新打包docx
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(temp_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, temp_dir)
                zipf.write(file_path, arcname)

    # 清理临时目录
    import shutil
    shutil.rmtree(temp_dir)

    return replaced


def fill_template(template_path, excel_data, output_path, debug=False):
    """
    填充Word模板
    """
    # 加载Word模板
    doc = Document(template_path)

    if not excel_data:
        print("警告：没有数据可供填充")
        return

    # 使用第一行数据（如需多行生成，需要多次处理）
    replacement_dict = excel_data[0]

    print(f"从Excel中加载了 {len(replacement_dict)} 个替换字段")
    if debug:
        print(f"替换字段列表: {list(replacement_dict.keys())}")

    # 检查是否是股东确认书
    template_name = os.path.basename(str(template_path))
    is_shareholder_confirmation = "股东确认书" in template_name
    
    # 处理股东确认书的特殊逻辑
    if is_shareholder_confirmation and "股东列表" in replacement_dict:
        try:
            shareholders = replacement_dict["股东列表"]
            print(f"处理股东确认书，股东数量: {len(shareholders)}")
            
            # 查找{{股东}}占位符的位置
            shareholder_placeholders = []
            
            # 检查表格中的{{股东}}
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            full_text = "".join([run.text for run in paragraph.runs])
                            placeholders = find_placeholders(full_text)
                            for placeholder_text, placeholder_name in placeholders:
                                if placeholder_name == "股东":
                                    shareholder_placeholders.append({
                                        "type": "table",
                                        "table": table,
                                        "table_idx": table_idx,
                                        "row_idx": row_idx,
                                        "cell_idx": cell_idx,
                                        "paragraph": paragraph
                                    })
            
            # 检查非表格中的{{股东}}
            for para_idx, paragraph in enumerate(doc.paragraphs):
                full_text = "".join([run.text for run in paragraph.runs])
                placeholders = find_placeholders(full_text)
                for placeholder_text, placeholder_name in placeholders:
                    if placeholder_name == "股东":
                        shareholder_placeholders.append({
                            "type": "paragraph",
                            "paragraph": paragraph,
                            "para_idx": para_idx
                        })
            
            print(f"找到 {len(shareholder_placeholders)} 个{{股东}}占位符")
            
            # 处理每个{{股东}}占位符
            for placeholder in shareholder_placeholders:
                if placeholder["type"] == "table":
                    # 表格中的{{股东}}
                    table = placeholder["table"]
                    row_idx = placeholder["row_idx"]
                    cell_idx = placeholder["cell_idx"]
                    paragraph = placeholder["paragraph"]
                    
                    if len(shareholders) > 1:
                        # 多个股东，替换第一个并添加行
                        # 替换第一个股东
                        first_shareholder = shareholders[0]
                        # 获取当前单元格并设置居中
                        current_cell = table.rows[row_idx].cells[cell_idx]
                        current_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        # 合并所有run的文本用于匹配
                        full_text = "".join([run.text for run in paragraph.runs])
                        placeholders_in_para = find_placeholders(full_text)
                        for placeholder_text, placeholder_name in placeholders_in_para:
                            if placeholder_name == "股东":
                                placeholder_pos = full_text.find(placeholder_text)
                                if placeholder_pos != -1:
                                    # 找到占位符跨越的所有run
                                    current_pos = 0
                                    runs_to_modify = []
                                    for run in paragraph.runs:
                                        run_start = current_pos
                                        run_end = current_pos + len(run.text)
                                        current_pos = run_end
                                        if run_start < placeholder_pos + len(placeholder_text) and run_end > placeholder_pos:
                                            runs_to_modify.append((run, run_start, run_end))
                                    
                                    if runs_to_modify:
                                        first_run, first_start, first_end = runs_to_modify[0]
                                        last_run, last_start, last_end = runs_to_modify[-1]
                                        
                                        if len(runs_to_modify) == 1:
                                            # 占位符完全在一个run中，直接替换
                                            first_run.text = first_run.text.replace(placeholder_text, first_shareholder)
                                        else:
                                            # 占位符跨越多个run，需要合并处理
                                            offset_in_first = placeholder_pos - first_start
                                            offset_in_last = (placeholder_pos + len(placeholder_text)) - last_start
                                            prefix = first_run.text[:offset_in_first] if offset_in_first > 0 else ""
                                            suffix = last_run.text[offset_in_last:] if offset_in_last < len(last_run.text) else ""
                                            for run, _, _ in runs_to_modify[1:-1]:
                                                run.text = ""
                                            first_run.text = prefix + first_shareholder + suffix
                                            last_run.text = ""
                                        # 设置段落水平居中
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 添加剩余股东的行
                    # 保存原始段落的格式信息
                    orig_paragraph = paragraph
                    orig_run_format = None
                    if orig_paragraph.runs:
                        orig_run_format = orig_paragraph.runs[0]
                    
                    for shareholder_idx, shareholder in enumerate(shareholders[1:], 1):
                        try:
                            # 复制当前行
                            new_row = table.add_row()
                            # 复制原行的格式
                            for cell_idx_copy, cell in enumerate(table.rows[row_idx].cells):
                                new_cell = new_row.cells[cell_idx_copy]
                                # 设置垂直居中 - 必须在添加内容之前设置
                                new_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                # 复制单元格内容和格式
                                for para_idx, cell_paragraph in enumerate(cell.paragraphs):
                                    new_paragraph = new_cell.add_paragraph()
                                    new_paragraph.style = cell_paragraph.style
                                    # 设置水平居中对齐
                                    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in cell_paragraph.runs:
                                        new_run = new_paragraph.add_run(run.text)
                                        new_run.bold = run.bold
                                        new_run.italic = run.italic
                                        new_run.underline = run.underline
                                        new_run.font.size = run.font.size
                                        new_run.font.name = run.font.name
                                        if run.font.color.rgb:
                                            new_run.font.color.rgb = run.font.color.rgb
                            # 替换新行中的股东占位符为当前股东
                            target_cell = new_row.cells[cell_idx]
                            # 清空单元格内容
                            target_cell.text = ""
                            # 重新设置垂直居中 - 在清空内容后再次设置
                            target_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            # 使用单元格中的第一个段落（清空后会有一个空段落）
                            if target_cell.paragraphs:
                                new_para = target_cell.paragraphs[0]
                            else:
                                new_para = target_cell.add_paragraph()
                            # 设置水平居中对齐
                            new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            new_run = new_para.add_run(shareholder)
                            # 复制原始段落的格式 - 确保所有字体属性都被复制
                            if orig_run_format:
                                # 复制字体名称
                                if orig_run_format.font.name:
                                    new_run.font.name = orig_run_format.font.name
                                # 复制字体大小
                                if orig_run_format.font.size:
                                    new_run.font.size = orig_run_format.font.size
                                # 复制粗体
                                new_run.bold = orig_run_format.bold
                                # 复制斜体
                                new_run.italic = orig_run_format.italic
                                # 复制下划线
                                new_run.underline = orig_run_format.underline
                                # 复制字体颜色
                                if orig_run_format.font.color and orig_run_format.font.color.rgb:
                                    new_run.font.color.rgb = orig_run_format.font.color.rgb
                                # 复制其他可能的字体属性
                                try:
                                    if orig_run_format.font.highlight_color:
                                        new_run.font.highlight_color = orig_run_format.font.highlight_color
                                except:
                                    pass
                        except Exception as e:
                            print(f"添加股东行时出错: {str(e)}")
                            continue
                elif placeholder["type"] == "table" and len(shareholders) == 1:
                    # 只有一个股东，直接替换
                    table = placeholder["table"]
                    row_idx = placeholder["row_idx"]
                    cell_idx = placeholder["cell_idx"]
                    paragraph = placeholder["paragraph"]
                    if shareholders:
                        # 获取当前单元格并设置居中
                        current_cell = table.rows[row_idx].cells[cell_idx]
                        current_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        # 合并所有run的文本用于匹配
                        full_text = "".join([run.text for run in paragraph.runs])
                        placeholders_in_para = find_placeholders(full_text)
                        for placeholder_text, placeholder_name in placeholders_in_para:
                            if placeholder_name == "股东":
                                placeholder_pos = full_text.find(placeholder_text)
                                if placeholder_pos != -1:
                                    # 找到占位符跨越的所有run
                                    current_pos = 0
                                    runs_to_modify = []
                                    for run in paragraph.runs:
                                        run_start = current_pos
                                        run_end = current_pos + len(run.text)
                                        current_pos = run_end
                                        if run_start < placeholder_pos + len(placeholder_text) and run_end > placeholder_pos:
                                            runs_to_modify.append((run, run_start, run_end))
                                    
                                    if runs_to_modify:
                                        first_run, first_start, first_end = runs_to_modify[0]
                                        last_run, last_start, last_end = runs_to_modify[-1]
                                        
                                        if len(runs_to_modify) == 1:
                                            # 占位符完全在一个run中，直接替换
                                            first_run.text = first_run.text.replace(placeholder_text, shareholders[0])
                                        else:
                                            # 占位符跨越多个run，需要合并处理
                                            offset_in_first = placeholder_pos - first_start
                                            offset_in_last = (placeholder_pos + len(placeholder_text)) - last_start
                                            prefix = first_run.text[:offset_in_first] if offset_in_first > 0 else ""
                                            suffix = last_run.text[offset_in_last:] if offset_in_last < len(last_run.text) else ""
                                            for run, _, _ in runs_to_modify[1:-1]:
                                                run.text = ""
                                            first_run.text = prefix + shareholders[0] + suffix
                                            last_run.text = ""
                                        # 设置段落水平居中
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # 非表格中的{{股东}}
                    paragraph = placeholder["paragraph"]
                    if len(shareholders) > 1:
                        # 多个股东用顿号拼接
                        shareholder_str = "、".join(shareholders)
                    else:
                        # 只有一个股东
                        shareholder_str = shareholders[0] if shareholders else ""
                    
                    # 替换占位符
                    full_text = "".join([run.text for run in paragraph.runs])
                    placeholders_in_para = find_placeholders(full_text)
                    for placeholder_text, placeholder_name in placeholders_in_para:
                        if placeholder_name == "股东":
                            placeholder_pos = full_text.find(placeholder_text)
                            if placeholder_pos != -1:
                                # 找到占位符跨越的所有run
                                current_pos = 0
                                runs_to_modify = []
                                for run in paragraph.runs:
                                    run_start = current_pos
                                    run_end = current_pos + len(run.text)
                                    current_pos = run_end
                                    if run_start < placeholder_pos + len(placeholder_text) and run_end > placeholder_pos:
                                        runs_to_modify.append((run, run_start, run_end))
                                
                                if runs_to_modify:
                                    first_run, first_start, first_end = runs_to_modify[0]
                                    last_run, last_start, last_end = runs_to_modify[-1]
                                    
                                    if len(runs_to_modify) == 1:
                                        # 占位符完全在一个run中，直接替换
                                        first_run.text = first_run.text.replace(placeholder_text, shareholder_str)
                                    else:
                                        # 占位符跨越多个run，需要合并处理
                                        offset_in_first = placeholder_pos - first_start
                                        offset_in_last = (placeholder_pos + len(placeholder_text)) - last_start
                                        prefix = first_run.text[:offset_in_first] if offset_in_first > 0 else ""
                                        suffix = last_run.text[offset_in_last:] if offset_in_last < len(last_run.text) else ""
                                        for run, _, _ in runs_to_modify[1:-1]:
                                            run.text = ""
                                        first_run.text = prefix + shareholder_str + suffix
                                        last_run.text = ""
        except Exception as e:
            print(f"处理股东确认书时出错: {str(e)}")
            import traceback
            traceback.print_exc()
    
    total_replaced = 0

    # 替换其他占位符
    if debug:
        print("\n=== 处理正文段落 ===")
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, replacement_dict, debug):
            total_replaced += 1

    # 替换表格中的其他占位符
    if debug:
        print("\n=== 处理正文表格 ===")
    for table in doc.tables:
        if replace_in_table(table, replacement_dict, debug):
            total_replaced += 1

    # 替换页眉中的占位符
    if replace_in_headers(doc, replacement_dict, debug):
        if debug:
            print("\n=== 页眉处理完成 ===")
        total_replaced += 1

    # 替换页脚中的占位符（python-docx方式）
    if replace_in_footers(doc, replacement_dict, debug):
        if debug:
            print("\n=== 页脚处理完成 ===")
        total_replaced += 1

    print(f"\n已完成正文替换，共处理 {total_replaced} 个段落/元素")

    # 保存中间结果
    temp_output = str(Path(output_path).parent / f"temp_{Path(output_path).name}")
    doc.save(temp_output)

    # 使用XML方式处理footer
    if replace_in_footers_xml(temp_output, replacement_dict, output_path, debug):
        print("Footer XML处理完成")
    else:
        # 如果XML处理没有替换，直接使用中间结果
        os.replace(temp_output, output_path)

    print(f"已保存到: {output_path}")


def process_all_templates(excel_path, debug=False):
    """
    处理所有Word模板文件
    1. 遍历脚本所在目录下所有.docx文件
    2. 为每个模板生成对应的输出文件
    3. 将所有生成的文档打包为ZIP压缩包
    """
    # 获取脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    print(f"脚本所在目录: {script_dir}")
    
    # 加载Excel数据
    excel_data = load_excel_data(excel_path)
    if not excel_data:
        print("错误：Excel中没有数据")
        return
    
    # 提取企业名称和额度启用日期用于ZIP命名
    data = excel_data[0]
    company_name = data.get("企业名称", "未命名企业")
    enable_date = data.get("额度启用日期", "未知日期")
    
    # 生成当前时间
    current_time = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    
    # 构造ZIP文件名：企业名称+额度启动日期+生成的此时时间
    zip_name = f"{company_name}_{enable_date}_{current_time}.zip"
    # 替换文件名中的非法字符
    zip_name = re.sub(r'[\\/:*?"<>|]', "_", zip_name)
    zip_output = os.path.join(script_dir, zip_name)

    # 查找所有.docx文件
    docx_files = []
    for file in os.listdir(script_dir):
        if file.endswith(".docx") and not file.startswith("~$"):  # 排除Office临时文件
            docx_files.append(os.path.join(script_dir, file))
    
    if not docx_files:
        print("错误：目录中没有找到.docx文件")
        return
    
    print(f"找到 {len(docx_files)} 个Word模板文件:")
    for docx_file in docx_files:
        print(f"  - {os.path.basename(docx_file)}")
    
    # 创建临时输出目录
    temp_output_dir = os.path.join(script_dir, "temp_output")
    os.makedirs(temp_output_dir, exist_ok=True)
    
    # 处理每个模板文件
    generated_files = []
    for template_path in docx_files:
        template_name = os.path.basename(template_path)
        output_file_name = template_name  # 保持与模板相同的名称
        output_path = os.path.join(temp_output_dir, output_file_name)
        
        print(f"\n处理模板: {template_name}")
        print(f"输出文件: {output_file_name}")
        
        # 执行填充
        fill_template(template_path, excel_data, output_path, debug=debug)
        generated_files.append(output_path)
    
    # 打包为ZIP文件
    print(f"\n打包所有生成的文档到: {zip_output}")
    
    with zipfile.ZipFile(zip_output, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file_path in generated_files:
            file_name = os.path.basename(file_path)
            zipf.write(file_path, file_name)
    
    print(f"\nZIP包创建成功！")
    print(f"生成的文档已打包到: {zip_output}")
    
    # 清理临时目录
    import shutil
    shutil.rmtree(temp_output_dir)


def main():
    parser = argparse.ArgumentParser(description="Word模板填充工具")
    parser.add_argument("--excel", default="input.xlsx", help="Excel数据文件路径（默认：input.xlsx）")
    parser.add_argument("--debug", action="store_true", help="显示详细调试信息")

    args = parser.parse_args()

    # 检查文件存在性
    excel_path = Path(args.excel)

    if not excel_path.exists():
        print(f"错误：Excel文件不存在: {excel_path}")
        sys.exit(1)

    print(f"=== Word模板填充工具 ===")
    print(f"Excel数据: {excel_path}")

    # 执行填充
    try:
        process_all_templates(excel_path, debug=args.debug)
        print("\n填充成功！")
    except Exception as e:
        print(f"\n填充失败: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
